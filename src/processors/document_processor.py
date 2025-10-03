"""Document text extraction processors."""

import os
from pathlib import Path
from typing import List, Optional, Tuple
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract

from src.processors.email_processor import EmailProcessor
from src.processors.excel_processor import ExcelProcessor
from src.utils import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Extract text from various document formats."""

    def __init__(self, enable_ocr: bool = True, tesseract_path: Optional[str] = None):
        """
        Initialize document processor.

        Args:
            enable_ocr: Enable OCR for scanned documents
            tesseract_path: Path to Tesseract executable
        """
        self.enable_ocr = enable_ocr
        self.email_processor = EmailProcessor()
        self.excel_processor = ExcelProcessor()

        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

    def process(self, file_path: str) -> List[Tuple[str, dict]]:
        """
        Process a document and extract text.

        Args:
            file_path: Path to document

        Returns:
            List of (text, metadata) tuples for each page/section
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path_obj.suffix.lower()

        logger.info(f"Processing document: {file_path_obj.name} (type: {extension})")

        if extension == ".pdf":
            return self._process_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return self._process_docx(file_path)
        elif extension in [".txt", ".text"]:
            return self._process_text(file_path)
        elif extension in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return self._process_image(file_path)
        elif extension in [".xlsx", ".xls"]:
            return self.excel_processor.process(file_path)
        elif extension == ".eml":
            return self.email_processor.process_eml(file_path)
        elif extension == ".msg":
            return self.email_processor.process_msg(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _process_pdf(self, file_path: str) -> List[Tuple[str, dict]]:
        """Extract text from PDF."""
        results = []

        try:
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Try direct text extraction first
                text = page.get_text()

                # If no text found and OCR enabled, use OCR
                if not text.strip() and self.enable_ocr:
                    logger.debug(f"No text found on page {page_num + 1}, attempting OCR")
                    text = self._ocr_pdf_page(page)

                if text.strip():
                    metadata = {
                        "page_number": page_num + 1,
                        "total_pages": len(doc),
                        "source": os.path.basename(file_path),
                        "extraction_method": "direct" if page.get_text().strip() else "ocr",
                    }
                    results.append((text, metadata))

            doc.close()
            logger.info(f"Extracted text from {len(results)} pages")

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

        return results

    def _ocr_pdf_page(self, page) -> str:
        """Perform OCR on a PDF page."""
        try:
            # Render page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Preprocess image
            img = self._preprocess_image(img)

            # Perform OCR
            text = pytesseract.image_to_string(img)
            return text

        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""

    def _process_docx(self, file_path: str) -> List[Tuple[str, dict]]:
        """Extract text from Word document."""
        results = []

        try:
            doc = Document(file_path)

            # Extract main document text
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            text = "\n".join(full_text)

            if text.strip():
                metadata = {
                    "source": os.path.basename(file_path),
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                }
                results.append((text, metadata))

            logger.info(f"Extracted {len(full_text)} text elements from DOCX")

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

        return results

    def _process_text(self, file_path: str) -> List[Tuple[str, dict]]:
        """Extract text from plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            metadata = {
                "source": os.path.basename(file_path),
                "encoding": "utf-8",
            }

            return [(text, metadata)]

        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()

            metadata = {
                "source": os.path.basename(file_path),
                "encoding": "latin-1",
            }

            return [(text, metadata)]

    def _process_image(self, file_path: str) -> List[Tuple[str, dict]]:
        """Extract text from image using OCR."""
        if not self.enable_ocr:
            logger.warning("OCR is disabled, cannot process image")
            return []

        try:
            img = Image.open(file_path)

            # Preprocess
            img = self._preprocess_image(img)

            # OCR
            text = pytesseract.image_to_string(img)

            metadata = {
                "source": os.path.basename(file_path),
                "width": img.width,
                "height": img.height,
                "extraction_method": "ocr",
            }

            return [(text, metadata)]

        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise

    def _preprocess_image(self, img: Image.Image) -> Image.Image:
        """
        Preprocess image for better OCR accuracy.

        Args:
            img: PIL Image

        Returns:
            Preprocessed image
        """
        # Convert to grayscale
        img = img.convert("L")

        # Increase contrast
        from PIL import ImageEnhance

        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)

        # Resize if too small (min 300 DPI equivalent)
        min_size = 1200
        if img.width < min_size or img.height < min_size:
            scale = max(min_size / img.width, min_size / img.height)
            new_size = (int(img.width * scale), int(img.height * scale))
            img = img.resize(new_size, Image.Resampling.LANCZOS)

        return img
